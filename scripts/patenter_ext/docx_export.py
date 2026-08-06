#!/usr/bin/env python3
"""M6: DOCX Executive Memo Export.

Adds native DOCX export for reports, complementing the existing Jinja2 HTML
pipeline. python-docx is imported *lazily* so the core patenter install stays
dependency-light — DOCX is an optional capability.

What it produces:
- A structured executive memo: title block, executive summary, stat table,
  findings sections, methodology note, disclaimer.
- Input is a JSON report context (same shape used by the render command's
  context-file), so it plugs into the existing data flow.

Usage (as a module):
    from modules.docx_export import export_docx_memo
    export_docx_memo(context_json_path="report.json", output="memo.docx")

CLI:
    python3 -m modules.docx_export report.json memo.docx
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

# Lazy import — only required when actually exporting DOCX
def _docx():
    try:
        import docx  # python-docx
        from docx.shared import Pt, RGBColor, Inches
        return docx, Pt, RGBColor, Inches
    except ImportError:
        raise RuntimeError(
            "python-docx not installed. Run: pip install python-docx"
        )


NAVY = None  # set after import


def _set_cell_bg(cell, hex_color: str) -> None:
    """Shade a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def export_docx_memo(context: dict, output_path: str | Path) -> Path:
    """Render a DOCX executive memo from a report context dict."""
    docx, Pt, RGBColor, Inches = _docx()
    doc = docx.Document()

    # --- base styles ---
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    navy = RGBColor(0x1F, 0x3A, 0x5F)

    # --- title block ---
    title = context.get("title", "Patent Intelligence Memo")
    doc.add_heading(title, level=0)
    if context.get("subtitle"):
        p = doc.add_paragraph()
        r = p.add_run(context["subtitle"])
        r.italic = True
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    meta = context.get("meta", {})
    if meta:
        doc.add_paragraph(f"Prepared: {meta.get('date', '')}   "
                          f"Technology: {meta.get('technology', '')}   "
                          f"Jurisdictions: {meta.get('jurisdictions', '')}")

    doc.add_paragraph("")  # spacer

    # --- executive summary ---
    if context.get("executive_summary"):
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(context["executive_summary"])

    # --- stat tiles as a table ---
    stats = context.get("stats", [])
    if stats:
        doc.add_heading("Key Metrics", level=1)
        table = doc.add_table(rows=2, cols=len(stats))
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        val = table.rows[1].cells
        for i, s in enumerate(stats[:8]):  # cap at 8 tiles
            hdr[i].text = str(s.get("label", ""))
            val[i].text = str(s.get("value", ""))
            _set_cell_bg(hdr[i], "1F3A5F")
            for r in hdr[i].paragraphs[0].runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True

    # --- findings sections ---
    for section in context.get("sections", []):
        doc.add_heading(section.get("title", "Findings"), level=1)
        for item in section.get("items", []):
            if isinstance(item, str):
                doc.add_paragraph(item, style="List Bullet")
            elif isinstance(item, dict):
                doc.add_paragraph(
                    f"{item.get('label', '')}: {item.get('value', '')}",
                    style="List Bullet",
                )

    # --- methodology ---
    if context.get("methodology"):
        doc.add_heading("Methodology", level=1)
        doc.add_paragraph(context["methodology"])

    # --- disclaimer ---
    if context.get("disclaimer"):
        p = doc.add_paragraph()
        r = p.add_run(context["disclaimer"])
        r.italic = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def export_docx_memo_from_file(context_path: str, output_path: str) -> Path:
    with open(context_path, "r", encoding="utf-8") as f:
        context = json.load(f)
    return export_docx_memo(context, output_path)


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Export a DOCX executive memo.")
    ap.add_argument("context", help="JSON report context file")
    ap.add_argument("output", help="Output .docx path")
    args = ap.parse_args()
    out = export_docx_memo_from_file(args.context, args.output)
    print(f"✅ DOCX memo saved: {out}")
